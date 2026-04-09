"""
resume_parser.py  —  Student Resume → Structured JSON Profile
=============================================================
Supports: PDF (via PyMuPDF/fitz) and DOCX (via python-docx)

Extracted fields:
  - cgpa           : float  (0–10)
  - projects       : int    (count)
  - internships    : int    (count)
  - communication  : int    (1–10, estimated)
  - open_source    : int    (0 / 1)
  - tech_stack     : list[str]
"""

import re
import io
from typing import Optional

# ── Tech dictionaries ──────────────────────────────────────────────────────────
STRONG_TECH = {
    "mern", "react", "angular", "vue", "next.js", "nextjs",
    "node", "node.js", "nodejs", "express", "fastapi", "django", "flask",
    "spring", "spring boot", "springboot",
    "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn",
    "machine learning", "deep learning", "nlp", "ai", "ml",
    "aws", "gcp", "azure", "kubernetes", "docker", "devops", "terraform",
    "go", "golang", "rust", "scala", "kotlin", "swift",
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "graphql", "grpc", "kafka", "rabbitmq",
}

MEDIUM_TECH = {
    "python", "java", "javascript", "typescript", "c++", "c#", "ruby",
    "php", "r", "matlab", "bash", "shell",
    "sql", "sqlite", "firebase", "supabase",
    "html", "css", "sass", "tailwind", "bootstrap",
    "git", "github", "gitlab", "jenkins", "ci/cd",
    "linux", "unix", "vim",
    "pandas", "numpy", "matplotlib", "seaborn",
    "selenium", "playwright", "jest", "pytest",
}

ALL_TECH = STRONG_TECH | MEDIUM_TECH

# Patterns for open-source detection
OSS_PATTERNS = [
    r"open[- ]?source", r"\bgithub\.com\b", r"\bpull request\b",
    r"\bpr(?:s)?\b", r"\bcontribut", r"\bopen source\b",
    r"\bfork\b", r"\brepository\b", r"\bmerged\b",
]

# Internship keywords
INTERN_PATTERNS = [
    r"\bintern\b", r"\binternship\b", r"\btrainee\b",
    r"\bwork(?:ed)? at\b", r"\bworked with\b", r"\bemployed\b",
    r"\bco-op\b", r"\bcoop\b", r"\bplacement\b",
    r"\bindustry experience\b", r"\bjunior developer\b",
]

# Project section headers
PROJECT_HEADERS = [
    r"projects?", r"personal projects?", r"academic projects?",
    r"key projects?", r"notable projects?", r"featured projects?",
    r"side projects?", r"open[- ]?source projects?",
]

# Communication quality signals
COMM_SIGNALS = {
    "high": [
        r"public speak", r"present", r"leadership", r"team lead",
        r"communicated", r"collaborated", r"mentor", r"negotiat",
        r"client", r"stakeholder", r"workshop", r"seminar",
    ],
    "medium": [
        r"teamwork", r"team player", r"coordinated", r"worked with",
        r"member", r"participate", r"discussion",
    ],
    "low": [],
}


# ── TEXT EXTRACTION ────────────────────────────────────────────────────────────

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        raise ValueError(f"PDF parsing failed: {e}")


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {e}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to correct parser based on filename extension."""
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return _extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return _extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext} (only PDF and DOCX supported)")


# ── FIELD EXTRACTORS ──────────────────────────────────────────────────────────

def _extract_cgpa(text: str) -> float:
    """Extract CGPA / GPA / percentage from text. Returns 7.0 as safe default."""
    # Pattern: CGPA: 8.5 / 10 or GPA: 3.8/4.0 or 8.5/10
    patterns = [
        r"(?:cgpa|gpa|cpi|spi)[:\s]*([0-9]+(?:\.[0-9]+)?)\s*/\s*([0-9]+(?:\.[0-9]+)?)",
        r"(?:cgpa|gpa|cpi|spi)[:\s]*([0-9]+(?:\.[0-9]+)?)",
        r"([0-9]+(?:\.[0-9]+)?)\s*/\s*10",
        r"([0-9]+(?:\.[0-9]+)?)\s*(?:cgpa|gpa|cpi)",
    ]
    text_lower = text.lower()
    for pat in patterns:
        matches = re.findall(pat, text_lower)
        if matches:
            m = matches[0]
            if isinstance(m, tuple):
                val = float(m[0])
                scale = float(m[1]) if m[1] else 10.0
                # Convert to 10-scale
                if scale == 4.0:
                    val = round((val / 4.0) * 10, 2)
                elif scale == 100.0 or val > 10:
                    val = round(val / 10, 2)
            else:
                val = float(m)
                if val > 10:
                    val = round(val / 10, 2)
            if 0.0 <= val <= 10.0:
                return round(val, 2)

    # Percentage fallback: 85% → 8.5
    pct_match = re.search(r"(\d{2,3}(?:\.\d+)?)\s*%", text_lower)
    if pct_match:
        pct = float(pct_match.group(1))
        if 40 <= pct <= 100:
            return round(pct / 10, 2)

    return 7.0  # safe default


def _extract_tech_stack(text: str) -> list:
    """Extract recognized technologies from resume text."""
    text_lower = text.lower()
    found = set()

    # Check for compound first (e.g. "next.js" before "next")
    for tech in sorted(ALL_TECH, key=len, reverse=True):
        # Word boundary match
        pattern = r'\b' + re.escape(tech) + r'\b'
        if re.search(pattern, text_lower):
            found.add(tech.title().replace(".Js", ".js").replace("Ai", "AI").replace("Ml", "ML"))

    # Additional: scan for lines that look like skill lists (comma/slash separated)
    skill_line = re.findall(
        r"(?:skills?|technologies?|tech stack|tools?|languages?)[:\s]+(.*?)(?:\n|$)",
        text_lower
    )
    for line in skill_line:
        tokens = re.split(r"[,|/•·\t]+", line)
        for token in tokens:
            token = token.strip()
            if 1 < len(token) < 30:
                for tech in ALL_TECH:
                    if tech in token:
                        found.add(tech.title())

    return sorted(list(found))


def _count_internships(text: str) -> int:
    """Count internship occurrences using heuristics."""
    text_lower = text.lower()
    count = 0

    # Check for experience section blocks
    # Split on section headers
    sections = re.split(
        r'\n(?:experience|work experience|professional experience|internships?)\s*\n',
        text_lower, flags=re.IGNORECASE
    )

    if len(sections) > 1:
        # Look in the experience section for company/role entries
        exp_section = sections[1] if len(sections) > 1 else ""
        # Each entry typically has a date range
        date_ranges = re.findall(
            r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*\d{4}"
            r"\s*[-–—to]+\s*"
            r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*\d{4}"
            r"|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*\d{4}"
            r"\s*[-–—to]+\s*(?:present|current|ongoing)",
            exp_section, flags=re.IGNORECASE
        )
        if date_ranges:
            count = len(date_ranges)

    # Fallback: keyword counting
    if count == 0:
        for pat in INTERN_PATTERNS:
            matches = re.findall(pat, text_lower)
            count += len(matches)
        count = min(count, 5)  # cap at 5

    # Check for explicit "X internship(s)"
    explicit = re.search(r"(\d+)\s+(?:internship|intern)\b", text_lower)
    if explicit:
        count = max(count, int(explicit.group(1)))

    return min(int(count), 6)


def _count_projects(text: str) -> int:
    """Count projects from resume text."""
    text_lower = text.lower()
    count = 0

    # Find project section and count bullet points / project names
    proj_section_match = re.search(
        r"(?:projects?|personal projects?|academic projects?|key projects?)"
        r"\s*\n(.*?)(?=\n(?:[A-Z][A-Za-z ]{2,20}\n)|$)",
        text, flags=re.DOTALL | re.IGNORECASE
    )

    if proj_section_match:
        proj_text = proj_section_match.group(1)
        # Count lines starting with bullet/number or title-case words (project names)
        bullets = re.findall(r"^[\s•\-\*\d\.]+\S", proj_text, re.MULTILINE)
        # Count GitHub links (each usually is a project)
        github_links = re.findall(r"github\.com/[^\s/]+/[^\s]+", text_lower)
        count = max(len(bullets), len(github_links))

    # Fallback: count "Project" word occurrences
    if count == 0:
        count = len(re.findall(r"\bproject\b", text_lower))
        count = min(count, 8)

    # Explicit "X projects" mention
    explicit = re.search(r"(\d+)\s+(?:projects?|applications?|systems?)\b", text_lower)
    if explicit:
        count = max(count, int(explicit.group(1)))

    return max(1, min(int(count), 10))


def _detect_open_source(text: str) -> int:
    """Return 1 if open-source contributions are evident, else 0."""
    text_lower = text.lower()
    matches = sum(
        1 for pat in OSS_PATTERNS
        if re.search(pat, text_lower)
    )
    # Strong signal: mentions GitHub with contributions
    has_github = bool(re.search(r"github\.com/[a-zA-Z0-9_-]+", text))
    has_contrib = bool(re.search(r"contribut|pull request|merged|fork", text_lower))
    if has_github and has_contrib:
        return 1
    return 1 if matches >= 2 else 0


def _estimate_communication(text: str, tech_stack: list) -> int:
    """
    Estimate communication score (1–10) from resume signals.
    Base: 5 (neutral)
    Boosts from leadership, presentation, teamwork keywords.
    """
    score = 5
    text_lower = text.lower()

    # High signals
    for pat in COMM_SIGNALS["high"]:
        if re.search(pat, text_lower):
            score += 0.8

    # Medium signals
    for pat in COMM_SIGNALS["medium"]:
        if re.search(pat, text_lower):
            score += 0.3

    # More tech → slightly higher (shows capacity to communicate complex topics)
    if len(tech_stack) >= 6:
        score += 0.5
    elif len(tech_stack) >= 3:
        score += 0.3

    # Resume structure is comprehensive → communication
    word_count = len(text.split())
    if word_count > 500:
        score += 0.5
    elif word_count > 300:
        score += 0.2

    return int(min(10, max(1, round(score))))


# ── MAIN PUBLIC API ───────────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Parse a student resume (PDF or DOCX) and return structured profile JSON.

    Returns:
    {
      "cgpa": float,
      "projects": int,
      "internships": int,
      "communication": int,
      "open_source": int,
      "tech_stack": list[str],
      "raw_text_length": int   (debug)
    }
    """
    text = extract_text(file_bytes, filename)

    tech_stack  = _extract_tech_stack(text)
    cgpa        = _extract_cgpa(text)
    projects    = _count_projects(text)
    internships = _count_internships(text)
    open_source = _detect_open_source(text)
    communication = _estimate_communication(text, tech_stack)

    return {
        "cgpa":          cgpa,
        "projects":      projects,
        "internships":   internships,
        "communication": communication,
        "open_source":   open_source,
        "tech_stack":    tech_stack,
        "raw_text_length": len(text),
    }


def parse_resume_from_text(text: str) -> dict:
    """Parse from already-extracted plain text (for testing/chat input)."""
    tech_stack    = _extract_tech_stack(text)
    cgpa          = _extract_cgpa(text)
    projects      = _count_projects(text)
    internships   = _count_internships(text)
    open_source   = _detect_open_source(text)
    communication = _estimate_communication(text, tech_stack)

    return {
        "cgpa":          cgpa,
        "projects":      projects,
        "internships":   internships,
        "communication": communication,
        "open_source":   open_source,
        "tech_stack":    tech_stack,
        "raw_text_length": len(text),
    }
